"""
Convert metbit_manuscript.md → metbit_manuscript.docx
Handles: headings, paragraphs, bold/italic/code inline, tables, images,
         horizontal rules, bullet lists, and blank-line paragraph breaks.
"""

import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent


# ── helpers ───────────────────────────────────────────────────────────────────

def set_run_font(run, name="Times New Roman", size=11, bold=False,
                 italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_horizontal_rule(doc):
    """Add a thin bottom-border paragraph as a visual divider."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_inline_runs(para, text, base_size=11, base_font="Times New Roman"):
    """
    Parse inline markdown (* ** ` [text](url)) and add styled runs.
    Handles: **bold**, *italic*, `code`, **bold *nested italic* bold**
    """
    # tokenise into segments: (text, bold, italic, code)
    # Strategy: walk character by character through a simple state machine
    segments = []
    i = 0
    bold = False
    italic = False
    code = False
    buf = []

    def flush():
        if buf:
            segments.append((''.join(buf), bold, italic, code))
            buf.clear()

    while i < len(text):
        # code span
        if text[i] == '`' and not code:
            flush()
            code = True
            i += 1
            continue
        if text[i] == '`' and code:
            flush()
            code = False
            i += 1
            continue

        # bold (**) — must check before single *
        if text[i:i+2] == '**' and not code:
            flush()
            bold = not bold
            i += 2
            continue

        # italic (*)
        if text[i] == '*' and not code:
            flush()
            italic = not italic
            i += 1
            continue

        # skip markdown image syntax entirely — handled at block level
        if text[i:i+2] == '![':
            j = text.find(')', i)
            i = j + 1 if j != -1 else i + 1
            continue

        # markdown link [label](url) → keep label only
        if text[i] == '[' and not code:
            end_bracket = text.find(']', i)
            if end_bracket != -1 and end_bracket + 1 < len(text) and text[end_bracket+1] == '(':
                end_paren = text.find(')', end_bracket)
                if end_paren != -1:
                    flush()
                    label = text[i+1:end_bracket]
                    buf.append(label)
                    flush()
                    i = end_paren + 1
                    continue
        buf.append(text[i])
        i += 1

    flush()

    for seg_text, is_bold, is_italic, is_code in segments:
        if not seg_text:
            continue
        run = para.add_run(seg_text)
        if is_code:
            run.font.name = "Courier New"
            run.font.size = Pt(base_size - 0.5)
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        else:
            run.font.name = base_font
            run.font.size = Pt(base_size)
            run.bold   = is_bold
            run.italic = is_italic


def parse_table(lines, doc):
    """Convert markdown table lines into a docx table."""
    # filter out separator rows (|---|---|)
    data_rows = [l for l in lines
                 if not re.match(r'^\s*\|[\s\-|:]+\|\s*$', l)]
    if not data_rows:
        return

    rows = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        rows.append(cells)

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= ncols:
                break
            cell = table.cell(r_idx, c_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

            is_header = (r_idx == 0)
            add_inline_runs(p, cell_text, base_size=9.5)
            if is_header:
                for run in p.runs:
                    run.bold = True
                # shade header row
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '1B4F72')
                tcPr.append(shd)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()   # spacing after table


def insert_image(doc, img_path_rel):
    """Try to insert an image; skip gracefully if not found."""
    img_path = HERE / img_path_rel
    if img_path.exists():
        try:
            doc.add_picture(str(img_path), width=Inches(5.8))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Image: {img_path_rel}]")
    else:
        doc.add_paragraph(f"[Image not found: {img_path_rel}]")


# ── document setup ────────────────────────────────────────────────────────────

def build_document(md_path: Path, out_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.17)
        section.right_margin  = Cm(3.17)

    # ── custom styles ─────────────────────────────────────────────────────────
    styles = doc.styles

    def _get_or_create(name, style_type=WD_STYLE_TYPE.PARAGRAPH):
        return styles[name] if name in [s.name for s in styles] \
               else styles.add_style(name, style_type)

    # Title style
    title_style = styles['Title']
    title_style.font.name = "Times New Roman"
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x15, 0x43, 0x60)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(10)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(13)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x15, 0x43, 0x60)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after  = Pt(4)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(11.5)
    h2.font.bold = True
    h2.font.italic = True
    h2.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after  = Pt(3)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x0E, 0x66, 0x55)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after  = Pt(2)

    # Normal body
    normal = styles['Normal']
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after  = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(0)

    # ── parse markdown ────────────────────────────────────────────────────────
    raw = md_path.read_text(encoding='utf-8')
    lines = raw.splitlines()

    # group consecutive table lines into blocks
    processed = []
    i = 0
    while i < len(lines):
        line = lines[i]
        # detect start of a table block
        if re.match(r'^\s*\|', line):
            block = []
            while i < len(lines) and re.match(r'^\s*\|', lines[i]):
                block.append(lines[i])
                i += 1
            processed.append(('TABLE', block))
        else:
            processed.append(('LINE', line))
            i += 1

    # ── render ────────────────────────────────────────────────────────────────
    for kind, payload in processed:
        if kind == 'TABLE':
            parse_table(payload, doc)
            continue

        line = payload

        # ── headings ──────────────────────────────────────────────────────────
        if line.startswith('#### '):
            p = doc.add_paragraph(line[5:].strip(), style='Heading 3')
            continue
        if line.startswith('### '):
            p = doc.add_paragraph(line[4:].strip(), style='Heading 3')
            continue
        if line.startswith('## '):
            p = doc.add_paragraph(line[3:].strip(), style='Heading 1')
            continue
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:].strip(), style='Title')
            continue

        # ── horizontal rule ───────────────────────────────────────────────────
        if re.match(r'^-{3,}\s*$', line):
            add_horizontal_rule(doc)
            continue

        # ── blank line → paragraph break (already implicit, skip) ─────────────
        if line.strip() == '':
            continue

        # ── image ─────────────────────────────────────────────────────────────
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', line)
        if img_match:
            insert_image(doc, img_match.group(2))
            continue

        # ── bullet list ───────────────────────────────────────────────────────
        bullet_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.75 + indent * 0.5)
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, bullet_match.group(2), base_size=10.5)
            continue

        # ── italic-only figure caption line (*Figure …*) ──────────────────────
        if re.match(r'^\*Figure', line) or re.match(r'^\*\*Figure', line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(10)
            add_inline_runs(p, line, base_size=9.5)
            continue

        # ── meta lines: Authors / affiliations / corresponding ────────────────
        if line.startswith('**Authors**'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, line, base_size=11)
            continue

        # ── general paragraph ─────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # affiliation / small-print lines (¹ ² * lines)
        if line.startswith(('¹', '²', '³', '\\*', '*Contributor')):
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, line.lstrip('\\'), base_size=9.5)
            continue

        # keywords line
        if line.startswith('**Keywords'):
            p.paragraph_format.space_before = Pt(4)
            add_inline_runs(p, line, base_size=10)
            continue

        add_inline_runs(p, line, base_size=11)

    doc.save(str(out_path))
    print(f"Saved → {out_path}")


# ── entry point ───────────────────────────────────────────────────────────────
if __name__ == '__main__':
    md   = HERE / 'metbit_manuscript.md'
    docx = HERE / 'metbit_manuscript.docx'
    build_document(md, docx)
